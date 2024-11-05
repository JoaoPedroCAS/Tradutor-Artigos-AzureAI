import requests
from docx import Document
import os

# Define a chave de assinatura e informações para a API do Microsoft Translator
subscription_key = "YOUR_KEY"
endpoint = "https://api.cognitive.microsofttranslator.com"
location = "YOUR_LOCATION"
target_language = "pt-br"

# Função para traduzir texto usando a API do Microsoft Translator
def translator_text(text, target_language):
    # Define o caminho do serviço de tradução (necessário configurar 'YOUR_PATH')
    path = "/translate"  # Ajustado para o caminho padrão da API de tradução
    constructed_url = endpoint + path  # Corrigido o erro de digitação em 'constructed_url'
    
    # Configuração do cabeçalho para a requisição da API
    headers = {
        'Ocp-Apim-Subscription-Key': subscription_key,
        'Ocp-Apim-Subscription-Region': location,
        'Content-type': 'application/json',
        'X-ClientTraceId': str(os.urandom(16))  # Identificador aleatório para a requisição
    }

    # Corpo da requisição com o texto a ser traduzido
    body = [{'text': text}]

    # Parâmetros da requisição com a versão da API e o idioma alvo
    params = {
        'api-version': '3.0',
        'from': 'en',  # Define o idioma de origem (inglês)
        'to': target_language  # Define o idioma alvo
    }

    # Faz a requisição à API e retorna a tradução
    response = requests.post(constructed_url, params=params, headers=headers, json=body)
    response_data = response.json()

    # Retorna o texto traduzido
    return response_data[0]["translations"][0]["text"]

# Função para traduzir o conteúdo de um documento DOCX
def translate_document(path):
    # Abre o documento de entrada
    document = Document(path)
    full_text = []  # Lista para armazenar os textos traduzidos

    # Itera sobre cada parágrafo do documento, traduzindo o texto e adicionando à lista
    for paragraph in document.paragraphs:
        if paragraph.text.strip():  # Ignora parágrafos vazios
            translated_text = translator_text(paragraph.text, target_language)
            full_text.append(translated_text)
        else:
            full_text.append("")  # Mantém parágrafos vazios

    # Cria um novo documento para salvar o texto traduzido
    translated_doc = Document()
    for line in full_text:
        translated_doc.add_paragraph(line)

    # Define o caminho para o documento traduzido e salva
    path_translated = path.replace(".docx", "_translated_document.docx")
    translated_doc.save(path_translated)
    return path_translated

# Define o caminho do arquivo de entrada e chama a função para traduzir o documento
input_file = "YOUR_INPUT_FILE"
translate_document(input_file)
